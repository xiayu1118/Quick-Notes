import re
import shutil
from datetime import datetime
import asyncio
from PyPDF2 import PdfReader
from flask import Blueprint, session,Response, stream_with_context
from langchain.embeddings.base import Embeddings
from sentence_transformers import SentenceTransformer
from typing import List
from langchain_openai import ChatOpenAI
import time
from waterproof import (remove_watermark,process_page,process_pdf,remove_watermark_from_image,process_image1,
                        process_pdf_and_extract_text)
import jwt
import json
from pymilvus import connections, Collection, utility, MilvusClient
import random
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from docx import Document
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_community.vectorstores import Milvus
bp = Blueprint('multimodels', __name__,url_prefix='/multimodels')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
api_key="ff18ba62b3fd3251ba7a7601db0992b8.9aXqtH8F6ona5kFb"

def generate_token(apikey: str, exp_seconds: int):
        try:
            id, secret = apikey.split(".")
        except Exception as e:
            raise Exception("invalid apikey", e)

        payload = {
            "api_key": id,
            "exp": int(round(time.time() * 1000)) + exp_seconds * 1000,
            "timestamp": int(round(time.time() * 1000)),
        }

        return jwt.encode(
            payload,
            secret,
            algorithm="HS256",
            headers={"alg": "HS256", "sign_type": "SIGN"},
        )

class SentenceTransformerEmbeddings(Embeddings):
    def __init__(self, model_name: str):
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, documents: List[str]) -> List[List[float]]:
        return [self.model.encode(d).tolist() for d in documents]

    def embed_query(self, query: str) -> List[float]:
        return self.model.encode([query])[0].tolist()
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 导入数据库模块
import pymysql
# 导入Flask框架，这个框架可以快捷地实现了一个WSGI应用
from config import attap_db
from flask import jsonify, request
import os
from milvus import (
    extract_text_from_pdf,
    extract_text_from_word, extract_text_from_txt, extract_text_from_image, split_text, load_file,
    extract_text_from_image1,
    document_to_dict, delete_collection
)
client=MilvusClient(
    uri="http://127.0.0.1:19530"
)
@bp.route("/show_filesss", methods=['POST'])
def show_filesss():
    #连接数据库
    conn = None
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')
    account=request.json.get('user')
    if not account:
        return jsonify({'success': 'false', 'message': '缺少用户名'}), 400
    sql = "SELECT * FROM words1 WHERE username = %s"
    cursor.execute(sql, (account,))
    wdnmd_data = list()
    wdnmd_data = cursor.fetchall()
    base_path = os.path.join('apps/static', 'milvus', account)
    conn.commit()
    cursor.close()
    conn.close()
    if not os.path.exists(base_path):
        os.makedirs(base_path)
    result = list()
    for ids in wdnmd_data:
         result.append({'id':ids[2],'name':ids[1],'imgSrc':'https://shadow.elemecdn.com/app/element/hamburger.9cf7b091-55e9-11e9-a976-7f4d0b07eef6.png'})
    print(result)
    return jsonify({'success': 'true', 'data': result}),200


@bp.route("/create_words", methods=['POST'])
def create_words():
    account = request.form.get("username")
    ids = request.form.get("ids")
    doc = request.files.get("doc")
    times = request.form.get("times")
    various = request.form.get("various")
    dconame = request.form.get("dconame")

    if not ids:
        return jsonify({'success': False, 'message': '缺少必要的参数 "ids"'}), 400

    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    cursor.execute('USE editdata')

    # 使用内置函数 id() 获取对象的唯一标识符，并将其转换为字符串
    idss1 = str(ids)

    # 拼接文件保存路径
    base_path = os.path.join('static', 'milvus', account, idss1)
    if not os.path.exists(base_path):
        os.makedirs(base_path)
    os.makedirs(base_path, exist_ok=True)
    text_parts = []
    # 如果上传了文件，则保存文件到对应文件夹中
    if doc:
        file_path = os.path.join(base_path, dconame)
        doc.save(file_path)
        # 插入数据库记录
        sql = "INSERT INTO words (dconame, various, times, username, id) VALUES (%s, %s, %s, %s, %s)"
        val = (dconame, various, times, account, idss1)
        cursor.execute(sql, val)
        conn.commit()
        ext = os.path.splitext(file_path)[1].lower()
        # 处理PDF文件，并将内容分割后存入Milvus
        documents = []
        pdf_reader=None
        if ext != '.pdf':
           pdf_reader = load_file(file_path)
        else:
           process_pdf(file_path,file_path)
           #使用pdfreader效果在去水印之后不好
           #pdf_reader= PdfReader(file_path)
           text_parts=process_pdf_and_extract_text(file_path,file_path)

        if ext != '.pdf':
            text_parts.append(pdf_reader)
        # 将每个部分的内容处理为Document对象
        for i, part in enumerate(text_parts):
            if len(part) > 65535:
                # 对超长文本进一步分割
                split_part = [part[j:j + 65535] for j in range(0, len(part), 65535)]
                for sub_part in split_part:
                    documents.append(Document(page_content=sub_part, metadata={"source": f"{dconame}_part_{i}"}))
            else:
                documents.append(Document(page_content=part, metadata={"source": f"{dconame}_part_{i}"}))

        embeddings = SentenceTransformerEmbeddings(r"C:\Users\lenovo\Desktop\others\EditEnd\chinese\chinese")

        # 将文档嵌入到 Milvus 中
        vector = Milvus.from_documents(
            documents=documents,  # 设置保存的文档
            embedding=embeddings,  # 设置 embedding model
            collection_name="items" + str(idss1),  # 设置集合名称
            drop_old=False,
            connection_args={"host": "127.0.0.1", "port": "19530"},  # Milvus连接配置
        )

        # 文档已成功嵌入到Milvus中
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'message': '文件保存成功，并已分割存入Milvus', 'path': file_path}), 200

    cursor.close()
    conn.close()
    return jsonify({'success': 'true', 'message': '文件夹已存在', 'path': base_path}), 200

@bp.route("/create_milvus", methods=['POST','GET'])
def create_milvus():
    account = request.form.get("username")
    ids = request.form.get("ids")
    kname=request.form.get("kname")
    conn = None
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')

    if not ids:
        return jsonify({'success': False, 'message': '缺少必要的参数 "ids"'}), 400
    idss1=str(ids)
    base_path = os.path.join('static', 'milvus', account, idss1)

    # 如果文件夹不存在，则创建文件夹
    if not os.path.exists(base_path):
        os.makedirs(base_path)
    # 如果文件存在并且上传了文件，则直接保存文件到对应文件夹中
    sql1 = "INSERT INTO words1 (username, kname, id, collections_name) VALUES (%s, %s, %s ,%s)"
    val1 = (str(account), str(kname), idss1 , idss1)
    try:
        cursor.execute(sql1, val1)
        conn.commit()
        return jsonify(message="Record inserted successfully."), 200
    except pymysql.MySQLError as e:
        return jsonify(error=str(e)), 500
    finally:
        cursor.close()
        conn.close()

@bp.route("/show_milvus", methods=['GET', 'POST'])
def show_milvus():
    #连接数据库
    conn = None
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')
    account = request.json.get("user")
    ids = request.json.get("id")
    idss1 = str(ids)
    if not account:
        return jsonify({'success': False, 'message': '缺少用户名'}), 400

    base_path = os.path.join('static', 'milvus', account)

    if not os.path.exists(base_path):
        os.makedirs(base_path)
    sql = "SELECT * FROM words WHERE username = %s AND id = %s"
    cursor.execute(sql, (account,idss1))
    wdnmd_data = list()
    wdnmd_data = cursor.fetchall()
    result =list()
    for ids in wdnmd_data:
        result.append({'date':ids[2],'name':ids[0],'sorts':ids[1]})
    conn.commit()
    return jsonify({'success': 'true', 'data': result})
class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

@bp.route("/show_session",methods=['GET', 'POST'])
def show_session():
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'), user=attap_db.get('USER'),
                           password=attap_db.get('PASSWORD'), charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('use editdata')
    account = request.json.get("username")
    sql="SELECT * FROM session WHERE account = %s"
    cursor.execute(sql, (account))
    wdnmd_data = list()
    wdnmd_data = cursor.fetchall()
    result = list()
    for ids in wdnmd_data:
        result.append(ids[1])
    return jsonify({'success': 'true', 'data': result})

@bp.route("/show_his", methods=['GET', 'POST'])
def show_his():
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'), user=attap_db.get('USER'),
                           password=attap_db.get('PASSWORD'), charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('use editdata')
    account = request.json.get("username")
    session_id=request.json.get("session_id")
    sql = "SELECT * FROM session WHERE account = %s AND session_id = %s"
    conn.commit()
    cursor.execute(sql, (account,session_id,))
    wdnmd_data = list()
    wdnmd_data = cursor.fetchall()
    result = list()
    for ids in wdnmd_data:
        messages = ids[2]
        # 将messages字符串转换为两个json对象
        messages_list = messages.split("}{")
        messages_list = [m + "}" if i == 0 else "{" + m for i, m in enumerate(messages_list)]
        a = 1
        for message in messages_list:
            # 使用正则表达式将单引号替换为双引号
            message = re.sub(r"'", r'"', message)
            message_dict = json.loads(message)
            new_format = {
                'id': a,
                'text': message_dict['content'],
                'sender': message_dict['role'],
                'avatar': 'userAvatar' if message_dict['role'] == 'user' else 'assistantAvatar'
            }
            result.append(new_format)

    print(result)
    return jsonify({'success': 'true', 'data': result})
@bp.route("/index_milvus", methods=['GET', 'POST'])
def index_milvus():
    account = request.form.get("username")
    idss = request.form.get("id")
    idss = str(idss)
    session_id=request.form.get("session_id")
    # 获取用户问题
    question = request.form.get("question")
    file_path = os.path.join('static', 'milvus', account, idss)
    #连接数据库
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'), user=attap_db.get('USER'),
                           password=attap_db.get('PASSWORD'), charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    embeddings = SentenceTransformerEmbeddings(r"C:\Users\lenovo\Desktop\others\EditEnd\chinese\chinese")
    #找出collections的id
    cursor.execute('use editdata')
    sql= "SELECT * FROM words1 WHERE username = %s AND id = %s"
    cursor.execute(sql, (account, idss,))
    conn.commit()

    collections_messages=cursor.fetchall()
    print(collections_messages)
    collection_id=collections_messages[0][3]
    print(collection_id)
    history =""
    # 创建或更新会话
    if session_id:
        #这个对话已经存在了
        sql = "SELECT * FROM session WHERE account = %s AND session_id = %s"
        cursor.execute(sql, (account, session_id))
        wdnmd_data = cursor.fetchall()
        for id in wdnmd_data:
            history+=id[2]
        conn.commit()
    else:
        #创建一个新的对话
        random_integer = random.randint(1, 60000)
        session_id=str(random_integer)
        sql="INSERT INTO session (account, session_id,history) VALUES (%s, %s,%s)"
        cursor.execute(sql, (account, str(random_integer),history))
        conn.commit()

    # 连接到 Milvus
    connections.connect("default", host="127.0.0.1", port="19530")

    # 定义集合名称
    collection_name = "items" + str(collection_id)

    # 初始化一个空列表来存储向量
    vector = None

    # 检查集合是否存在
    if utility.has_collection(collection_name):
        vector = Milvus(
            embeddings,
            connection_args={"host": "127.0.0.1", "port": "19530"},
            collection_name=collection_name,
        )
    else:
        documents = []
        text1 = ""
        file_full_path = file_path
        for dirpath, dirnames, filenames in os.walk(file_path):
            for filename in filenames:
                file_full_path = os.path.join(dirpath, filename)
                if os.path.isfile(file_full_path):  # 判断是否是文件
                    text = load_file(file_full_path)
                    text1 += text
        text1 += str(history)
        documents.append(Document(page_content=text1, metadata={"source": ""}))
        # 将文档嵌入到 Milvus 中
        vector = Milvus.from_documents(
            documents=documents,  # 设置保存的文档
            embedding=embeddings,  # 设置 embedding model
            collection_name=collection_name,  # 设置 集合名称
            drop_old=True,
            connection_args={"host": "127.0.0.1", "port": "19530"},  # Milvus连接配置
        )
    # Prompt 模板
    prompt = ChatPromptTemplate.from_template("""You are an assistant for question-answering tasks. Use the following pieces of retrieved context to answer the question. If you don't know the answer, just say that you don't know.
    Your answer is best presented in the form of a table of contents, without the need for superfluous words, such as:
    An internal overview of what the user needs to retrieve:
    1.XXX
    2.XXX
    3.XXX

    Question: {question}

    Context: {context}

    Answer:
    """)

    # 加载 Chat 模型
    token_key = "ff18ba62b3fd3251ba7a7601db0992b8.9aXqtH8F6ona5kFb"
    llm = ChatOpenAI(
        model_name="glm-4",
        openai_api_base="https://open.bigmodel.cn/api/paas/v4",
        openai_api_key=generate_token(token_key, exp_seconds=2592000),
        streaming=True,
        temperature=0.01
    )

    retriever = vector.as_retriever()
    # 将新问题添加到历史上下文
    history+=str(({"role": "user", "content": question}))
    # 执行检索并存储结果
    context_results = vector.as_retriever()
    # 假设 context_results 的结构是 [(text1, vector1), (text2, vector2), ...]
    context_texts = [result[0] for result in context_results]  # 访问第一个元素（文本）
    context = "\n".join(context_texts)
    # 将检索结果添加到历史上下文
    wdnmd=({"role": "retriever", "content": context})
    # 使用上下文生成答案
    chain = (
        RunnableParallel({"context": retriever, "question": RunnablePassthrough()})
        | prompt
        | llm
        | StrOutputParser()
    )

    def gen():
        for chunk in chain.stream({"context": context, "question": question}):
            yield chunk

    answer = chain.invoke({"context": context, "question": question})
    # 将答案添加到历史上下文
    history+=str(({"role": "assistant", "content": answer}))
    sql = "UPDATE session SET history = %s WHERE account = %s AND session_id = %s"
    cursor.execute(sql, (history, account, session_id))
    conn.commit()
    # 保存更新的会话历史
    #return Response(stream_with_context(gen()), content_type='text/plain')
    return jsonify({'success': 'true', 'answer': answer,'session_id':str(session_id),'history':str(history),'retriever':str(retriever)})
@bp.route('/delete_session', methods=['GET', 'POST'])
def delete_session():
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'), user=attap_db.get('USER'),
                           password=attap_db.get('PASSWORD'), charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('use editdata')
    session_id = request.form.get("session_id")
    account=request.form.get("username")
    sql = "DELETE FROM session WHERE account = %s AND session_id = %s"
    cursor.execute(sql, (account, session_id))
    conn.commit()

@bp.route('/get_chunks', methods=['POST', 'GET'])
def get_chunks():
    account = request.form.get("username")
    idss = request.form.get("id")
    dconame = request.form.get("dconame")
    file_path = os.path.join(r'C:\Users\lenovo\Desktop\others\EditEnd\apps\static', 'milvus', account, idss, dconame)

    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".txt":
        text = extract_text_from_txt(file_path)
    elif ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
        text = extract_text_from_image1(file_path)
    elif ext == ".docx":
        text = extract_text_from_word(file_path)
        chunks_dict =""
        for a in text:
            chunks_dict+=a.page_content
        text=chunks_dict
    if ext == ".pdf":
      # 将文本分割为适当大小的块
      chunks = split_text(text["text"])
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
        chunks=text.content
        chunks = split_text(chunks)
    else:
      chunks = split_text(text)
    # 构建文件 URL（假设文件可以通过 /static/ 路径访问）
    base_url = request.host_url  # e.g., "http://localhost:5000/"
    file_url = os.path.join(base_url, 'static', 'milvus', account, idss, dconame).replace("\\", "/")
    print(file_url)
    return jsonify({'success': 'true', 'chunk': chunks, 'file_url': file_url})
@bp.route('/delete_file', methods=['GET', 'POST'])
def delete_file():
    account = request.form.get("username")
    idss = request.form.get("ids")
    idss = str(idss)
    dconame = request.form.get("dconame")
    file_path = os.path.join(r'C:\Users\lenovo\Desktop\others\EditEnd\apps\static', 'milvus', account, idss, dconame)
    ext = os.path.splitext(file_path)[1].lower()
    conn = None
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')
    sql="DELETE FROM words WHERE username = %s AND dconame = %s AND id = %s"
    cursor.execute(sql, (account, dconame ,idss))
    conn.commit()
    cursor.execute('SELECT * FROM words1 WHERE id = %s AND username = %s', (idss,account,))
    conn.commit()
    result = cursor.fetchone()
    collections_name = str(result[3])
    collections_name=("items"+collections_name)
    source_to_delete = file_path
    # 转义反斜杠
    source_to_delete_escaped = source_to_delete.replace("\\", "\\\\")
    # 构造删除的过滤表达式
    filter_expr = f"source == '{source_to_delete_escaped}'"
    # 删除milvus数据库中的向量
    print(collections_name)
    res = client.delete(
        collection_name=collections_name,
        filter=filter_expr
    )
    print(res)
    if os.path.exists(file_path):
        # 删除文件
        os.remove(file_path)
    return jsonify({'success': 'true','message':'deleted successfully'})

@bp.route('/delete_word', methods=['GET', 'POST'])
def delete_word():
    account = request.form.get("username")
    idss = request.form.get("ids")
    #知识库整个删除
    idss = str(idss)
    conn = None
    conn = pymysql.connect(host=attap_db.get('HOST'), port=attap_db.get('PORT'),
                           user=attap_db.get('USER'), password=attap_db.get('PASSWORD'),
                           charset=attap_db.get('CHARSET'))
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')
    sql1 = "SELECT * FROM words1 WHERE username = %s AND id = %s"
    cursor.execute(sql1, (account,idss))
    conn.commit()
    temp=cursor.fetchall()
    print(temp)
    collections_name=temp[0][3]
    collections_name="items"+str(collections_name)
    #取出向量数据库的collection的名字
    conn.commit()
    cursor = conn.cursor()
    conn.commit()
    cursor.execute('USE editdata')
    sql2 = "DELETE FROM words1 WHERE username = %s AND id = %s"
    cursor.execute(sql2, (account, idss))
    conn.commit()
    sql3 = "DELETE FROM words WHERE username = %s AND id = %s"
    cursor.execute(sql3, (account, idss))
    conn.commit()
    file_path = os.path.join(r'C:\Users\lenovo\Desktop\others\EditEnd\apps\static', 'milvus', account, idss)
    try:
        shutil.rmtree(file_path)
    except Exception as e:
        print(f"发生错误: {e}")
    #把对应的milvus向量知识库代码删除
    # 连接到 Milvus
    connections.connect("default", host="127.0.0.1", port="19530")
    # 指定要删除的集合的名字
    collection_name = collections_name
    # 实例化 Collection 对象
    if utility.has_collection(collection_name):
        # 实例化 Collection 对象
        collection = Collection(name=collection_name)
        # 删除集合
        collection.drop()
    return jsonify({'success': 'true','message':'deleted successfully'})



